# -*- coding: utf-8 -*-
"""Generate Hikvision-Odoo bridge documentation using company Word template."""

from __future__ import annotations

import shutil
import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

TEMPLATE = Path(r"c:\Users\ASUS\Downloads\ملف قالب الشركة (2).docx")
OUTPUT_DOCX = Path(__file__).resolve().parent / "Hikvision_Odoo_Attendance_Bridge.docx"
OUTPUT_PDF = Path(__file__).resolve().parent / "Hikvision_Odoo_Attendance_Bridge.pdf"
DOWNLOADS_PDF = Path(r"c:\Users\ASUS\Downloads\Hikvision_Odoo_Attendance_Bridge.pdf")

COLOR_PRIMARY = RGBColor(0x0E, 0x28, 0x41)
COLOR_ACCENT = RGBColor(0x15, 0x60, 0x82)
COLOR_ACCENT2 = RGBColor(0xE9, 0x71, 0x32)
COLOR_TEXT = RGBColor(0x33, 0x33, 0x33)


def set_paragraph_shading(paragraph, fill_hex: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    p_pr.append(shd)


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Hikvision Fingerprint to Odoo Attendance")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = COLOR_PRIMARY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Integration Guide — Option B (FastAPI Bridge)")
    r.font.size = Pt(16)
    r.font.color.rgb = COLOR_ACCENT

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        f"Odoo: https://shamsieh-testing-33874628.dev.odoo.com\n"
        f"Webhook: http://192.168.100.4:8080/hikvision/attendance\n"
        f"Document date: {date.today().strftime('%B %d, %Y')}"
    )
    r.font.size = Pt(11)
    r.font.color.rgb = COLOR_TEXT
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = COLOR_PRIMARY
        set_paragraph_shading(p, "E8E8E8")
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = COLOR_ACCENT
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_ACCENT2


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_TEXT


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"• {text}")
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_TEXT
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
            set_paragraph_shading(paragraph, "0E2841")
    for row_index, row in enumerate(rows):
        cells = table.rows[row_index + 1].cells
        for col_index, value in enumerate(row):
            cells[col_index].text = str(value)
            for paragraph in cells[col_index].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()


def build_document() -> Document:
    if TEMPLATE.exists():
        shutil.copy2(TEMPLATE, OUTPUT_DOCX)
        doc = Document(str(OUTPUT_DOCX))
        body = doc.element.body
        for child in list(body):
            if child.tag.endswith("p"):
                body.remove(child)
    else:
        doc = Document()

    add_cover(doc)

    add_heading(doc, "1. Overview", 1)
    add_body(
        doc,
        "This integration connects a Hikvision fingerprint terminal on the local network to "
        "Odoo Attendances on Odoo.sh. A FastAPI bridge runs on a Windows PC, receives device "
        "events, and creates hr.attendance check-ins through XML-RPC. No Fingerprint Device "
        "configuration is required inside Odoo."
    )

    add_heading(doc, "2. Architecture", 1)
    add_table(
        doc,
        ["Component", "Role"],
        [
            ["Hikvision device (192.168.100.85)", "Captures fingerprint; pushes HTTP events"],
            ["FastAPI bridge (192.168.100.4:8080)", "Parses events, idempotency, retries"],
            ["Odoo.sh", "Stores employees and attendance records"],
            ["SQLite (local)", "Processed events + retry queue"],
        ],
    )
    add_body(doc, "Flow: Device POST → /hikvision/attendance → find employee by barcode → create check-in.")

    add_heading(doc, "3. Configuration", 1)
    add_heading(doc, "3.1 Hikvision HTTP Listening", 2)
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Event Alarm IP/Domain Name", "192.168.100.4"],
            ["Port", "8080"],
            ["URL", "/hikvision/attendance"],
            ["Protocol", "HTTP"],
        ],
    )
    add_heading(doc, "3.2 Odoo Employee Mapping", 2)
    add_bullet(doc, "RFID/Badge Number = Hikvision employeeNoString (example: 5)")
    add_bullet(doc, "Biometric Device User ID = same value (recommended)")

    add_heading(doc, "3.3 Bridge Environment (.env)", 2)
    add_table(
        doc,
        ["Variable", "Example"],
        [
            ["ODOO_URL", "https://shamsieh-testing-33874628.dev.odoo.com"],
            ["ODOO_DB", "shamsieh-testing-33874628"],
            ["ODOO_BOT_USER", "admin"],
            ["ODOO_API_KEY", "API key from Odoo Preferences"],
            ["LISTEN_PORT", "8080"],
        ],
    )

    add_heading(doc, "4. HTTP API & Status Codes", 1)
    add_body(
        doc,
        "The bridge returns meaningful HTTP status codes so operators and integrators can "
        "distinguish success, wrong fingerprint, missing employee, and system events."
    )

    add_heading(doc, "4.1 Endpoints", 2)
    add_table(
        doc,
        ["Method", "URL", "Purpose"],
        [
            ["GET", "/health", "Service health check"],
            ["GET", "/odoo/ping", "Test Odoo.sh XML-RPC connection"],
            ["GET", "/hikvision/attendance", "Webhook reachability check"],
            ["POST", "/hikvision/attendance", "Receive Hikvision fingerprint events"],
        ],
    )

    add_heading(doc, "4.2 POST /hikvision/attendance — Status Codes", 2)
    add_table(
        doc,
        ["HTTP", "reason / result", "When it happens"],
        [
            ["201", "created", "Fingerprint success; new hr.attendance check-in created"],
            ["200", "duplicate-attendance", "Employee already checked in today"],
            ["200", "duplicate-event", "Same device event already processed (idempotency)"],
            ["404", "employee-not-found", "employeeNoString not found in Odoo barcode"],
            ["422", "fingerprint-failed", "Wrong finger / failed scan (subEventType 39)"],
            ["422", "missing-employee", "Event has no employee number"],
            ["422", "not-fingerprint-event", "Not a fingerprint authentication event"],
            ["204", "system-event", "Door/system event (subEventType 21–24); ignored"],
            ["400", "invalid-payload", "Malformed or unparseable request body"],
            ["503", "odoo-unavailable", "Odoo.sh unreachable; event queued for retry"],
        ],
    )

    add_heading(doc, "4.3 Hikvision Event Codes", 2)
    add_table(
        doc,
        ["subEventType", "Meaning", "Bridge action"],
        [
            ["38", "Fingerprint verification success", "Process attendance if employee mapped"],
            ["39", "Fingerprint verification failed", "HTTP 422 fingerprint-failed"],
            ["21", "Door locked", "HTTP 204 system-event (ignored)"],
            ["22", "Door unlocked", "HTTP 204 system-event (ignored)"],
        ],
    )

    add_heading(doc, "5. Postman Test Bodies", 1)
    add_heading(doc, "5.1 Success (expect 201)", 2)
    add_body(
        doc,
        'POST http://192.168.100.4:8080/hikvision/attendance\n'
        'Content-Type: application/json\n\n'
        '{"subEventType":38,"employeeNoString":"5","currentVerifyMode":"fp",'
        '"dateTime":"2026-07-07T13:40:51+03:00","serialNo":9999,"status":"success"}'
    )
    add_heading(doc, "5.2 Wrong fingerprint (expect 422)", 2)
    add_body(
        doc,
        'POST http://192.168.100.4:8080/hikvision/attendance\n'
        'Content-Type: application/json\n\n'
        '{"subEventType":39,"currentVerifyMode":"faceOrFpOrCardOrPw","status":"failed","serialNo":9998}'
    )

    add_heading(doc, "6. Service Features", 1)
    for item in [
        "Parses Hikvision multipart/form-data JSON (event_log field).",
        "Uses webhook receive time when device omits dateTime.",
        "SQLite idempotency on (device_serial, event_id).",
        "Retries failed Odoo calls every 30 seconds.",
        "Auto-start via Windows scheduled task HikvisionAttendanceBridge.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "7. Start & Verify", 1)
    add_bullet(doc, "Manual: .\\run.ps1 in scripts/hikvision_attendance_service")
    add_bullet(doc, "Health: GET http://192.168.100.4:8080/health")
    add_bullet(doc, "Odoo: GET http://192.168.100.4:8080/odoo/ping")
    add_bullet(doc, "Fingerprint on device → check Attendances in Odoo.sh")

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("— End of Document —")
    run.font.color.rgb = COLOR_ACCENT
    run.italic = True
    return doc


def convert_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        return pdf_path.exists()
    except Exception as exc:
        print(f"docx2pdf failed: {exc}")
    try:
        soffice = Path(r"C:\Program Files\LibreOffice\program\soffice.exe")
        if soffice.is_file():
            subprocess.run(
                [str(soffice), "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)],
                check=True,
                timeout=120,
            )
            return pdf_path.exists()
    except Exception as exc:
        print(f"LibreOffice failed: {exc}")
    return False


def main() -> None:
    print("Building documentation from company template...")
    doc = build_document()
    doc.save(str(OUTPUT_DOCX))
    print(f"DOCX: {OUTPUT_DOCX}")

    if convert_to_pdf(OUTPUT_DOCX, OUTPUT_PDF):
        shutil.copy2(OUTPUT_PDF, DOWNLOADS_PDF)
        print(f"PDF: {OUTPUT_PDF}")
        print(f"Copied to: {DOWNLOADS_PDF}")
    else:
        print("PDF conversion failed. Open the DOCX in Word and Save As PDF.")


if __name__ == "__main__":
    main()
