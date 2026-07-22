# -*- coding: utf-8 -*-
"""Generate simple Hikvision-Odoo bridge guide (company Word template)."""

from __future__ import annotations

import shutil
import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

TEMPLATE = Path(r"c:\Users\ASUS\Downloads\ملف قالب الشركة (2).docx")
OUTPUT_DOCX = Path(__file__).resolve().parent / "Hikvision_Odoo_Attendance_Bridge.docx"
OUTPUT_PDF = Path(__file__).resolve().parent / "Hikvision_Odoo_Attendance_Bridge.pdf"
DOWNLOADS_PDF = Path(r"c:\Users\ASUS\Downloads\Hikvision_Odoo_Attendance_Bridge.pdf")

COLOR_PRIMARY = RGBColor(0x0E, 0x28, 0x41)
COLOR_ACCENT = RGBColor(0x15, 0x60, 0x82)
COLOR_ACCENT2 = RGBColor(0xE9, 0x71, 0x32)
COLOR_TEXT = RGBColor(0x33, 0x33, 0x33)


def set_paragraph_shading(paragraph, fill_hex: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    p_pr.append(shd)


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Hikvision → Odoo Attendance")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = COLOR_PRIMARY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Quick Setup Guide (read this when you come back)")
    r.font.size = Pt(16)
    r.font.color.rgb = COLOR_ACCENT

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        f"Updated: {date.today().strftime('%B %d, %Y')}\n"
        "Folder: scripts\\hikvision_attendance_service"
    )
    r.font.size = Pt(11)
    r.font.color.rgb = COLOR_TEXT
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = COLOR_PRIMARY
        set_paragraph_shading(p, "E8E8E8")
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = COLOR_ACCENT
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_ACCENT2


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_TEXT


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"• {text}")
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_TEXT
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
            set_paragraph_shading(paragraph, "0E2841")
    for row_index, row in enumerate(rows):
        cells = table.rows[row_index + 1].cells
        for col_index, value in enumerate(row):
            cells[col_index].text = str(value)
            for paragraph in cells[col_index].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()


def build_document() -> Document:
    if TEMPLATE.exists():
        shutil.copy2(TEMPLATE, OUTPUT_DOCX)
        doc = Document(str(OUTPUT_DOCX))
        body = doc.element.body
        for child in list(body):
            if child.tag.endswith("p"):
                body.remove(child)
    else:
        doc = Document()

    add_cover(doc)

    add_heading(doc, "What is this?", 1)
    add_body(
        doc,
        "A small Python program on your Windows PC receives fingerprint events from the Hikvision "
        "device and creates check-ins in Odoo Attendances on Odoo.sh. "
        "You do NOT configure anything inside Odoo for the fingerprint device itself."
    )
    add_body(doc, "Flow: Fingerprint on device → PC bridge → Odoo.sh attendance record.")

    add_heading(doc, "Your machines (write these down)", 1)
    add_table(
        doc,
        ["What", "Address"],
        [
            ["Hikvision device", "192.168.100.85"],
            ["Bridge PC (this PC)", "192.168.100.4"],
            ["Bridge URL", "http://192.168.100.4:8080/hikvision/attendance"],
            ["Odoo.sh", "https://shamsieh-testing-33874628.dev.odoo.com"],
            ["Odoo database name", "shamsieh-testing-33874628"],
        ],
    )

    add_heading(doc, "Start the bridge (every time)", 1)
    add_body(doc, "Open PowerShell:")
    add_body(
        doc,
        "cd C:\\Users\\ASUS\\Desktop\\odoo\\scripts\\hikvision_attendance_service\n"
        ".\\.venv\\Scripts\\Activate.ps1\n"
        "uvicorn app.main:app --host 0.0.0.0 --port 8080"
    )
    add_body(doc, "Leave the window open. Test in browser:")
    add_bullet(doc, "http://127.0.0.1:8080/health  →  {\"status\":\"ok\"}")
    add_bullet(doc, "http://127.0.0.1:8080/odoo/ping  →  {\"status\":\"ok\"}")

    add_heading(doc, "Auto-start on Windows login", 1)
    add_body(doc, "Already installed to:")
    add_bullet(doc, "C:\\Users\\ASUS\\AppData\\Local\\HikvisionAttendanceBridge")
    add_bullet(doc, "Scheduled task name: HikvisionAttendanceBridge")
    add_body(doc, "Restart task: Start-ScheduledTask -TaskName HikvisionAttendanceBridge")

    add_heading(doc, "Hikvision device settings", 1)
    add_body(doc, "Configuration → Network → Advanced → HTTP Listening:")
    add_table(
        doc,
        ["Field", "Value"],
        [
            ["Event Alarm IP", "192.168.100.4"],
            ["Port", "8080"],
            ["URL", "/hikvision/attendance"],
            ["Protocol", "HTTP"],
        ],
    )
    add_body(
        doc,
        "Important: The device must send authentication events, not only door status. "
        "If you only see subEventType 21–24 in logs, enable access/authentication events in the device event settings."
    )

    add_heading(doc, "Odoo employee setup (required)", 1)
    add_body(doc, "For each person on the fingerprint device:")
    add_bullet(doc, "Open Employees → select person")
    add_bullet(doc, "RFID/Badge Number = same number as on Hikvision (example: 5 for MohammadNoor)")
    add_bullet(doc, "Biometric Device User ID = same number (recommended)")
    add_body(
        doc,
        "Example verified: MohammadNoor (employee id 3) with barcode 5 matches device employeeNoString 5."
    )

    add_heading(doc, "Odoo login for the bridge (.env file)", 1)
    add_body(doc, "The bridge uses XML-RPC + API key. It does NOT use PostgreSQL.")
    add_body(doc, "File: scripts\\hikvision_attendance_service\\.env")
    add_table(
        doc,
        ["Variable", "Value"],
        [
            ["ODOO_URL", "https://shamsieh-testing-33874628.dev.odoo.com"],
            ["ODOO_DB", "shamsieh-testing-33874628"],
            ["ODOO_BOT_USER", "Your Odoo login email (e.g. m.saqer@shamsieh.com)"],
            ["ODOO_API_KEY", "From Odoo → Preferences → Account Security → New API Key"],
            ["LISTEN_PORT", "8080"],
            ["VERBOSE_LOGGING", "true = see every device event in console"],
        ],
    )
    add_body(
        doc,
        "If /odoo/ping fails: wrong email or API key. Generate a new API key while logged in as that user."
    )

    add_heading(doc, "What the device sends (event types)", 1)
    add_table(
        doc,
        ["subEventType", "Meaning", "What happens"],
        [
            ["38 or 150", "Fingerprint / auth SUCCESS", "Creates attendance in Odoo (if barcode matches)"],
            ["39 or 151", "Fingerprint / auth FAILED", "Logged as fingerprint-failed; no Odoo record"],
            ["21–24", "Door / system noise", "Ignored (door locked, unlocked, etc.)"],
        ],
    )
    add_body(doc, "Success log looks like:")
    add_body(doc, "employee_no='5'  subEventType=38  statusValue=1  →  Created attendance id=...")
    add_body(doc, "Failed log looks like:")
    add_body(doc, "employee_no=None  subEventType=151  statusValue=0  →  wrong finger or not enrolled on device")

    add_heading(doc, "Webhook responses (JSON body)", 1)
    add_body(
        doc,
        "The device always gets HTTP 200 or 201 (so it stops retrying). "
        "Read the JSON reason field to know what happened:"
    )
    add_table(
        doc,
        ["reason / result", "Meaning"],
        [
            ["created (HTTP 201)", "Check-in created in Odoo"],
            ["duplicate-attendance", "Already checked in today"],
            ["duplicate-event", "Same event already processed"],
            ["employee-not-found", "Barcode not set in Odoo for that employee number"],
            ["fingerprint-failed", "Device rejected the scan"],
            ["system-event", "Door event — ignored"],
            ["odoo-unavailable", "Odoo down — queued for retry"],
        ],
    )

    add_heading(doc, "Troubleshooting cheat sheet", 1)
    add_table(
        doc,
        ["Problem", "Fix"],
        [
            ["Only subEventType 21–24 in logs", "Device not sending auth events — check Hikvision event settings"],
            ["subEventType 151, no employee_no", "Fingerprint failed on device — re-enroll finger"],
            ["employee-not-found", "Set RFID/Badge Number in Odoo = device employee ID"],
            ["Odoo connection failed", "Fix ODOO_BOT_USER (email) and ODOO_API_KEY in .env"],
            ["Same event repeats forever", "Restart bridge (latest code returns 200 to ACK events)"],
            ["Nothing in Odoo server logs", "Normal — bridge talks via XML-RPC from your PC, not Odoo web"],
        ],
    )

    add_heading(doc, "Test without the device", 1)
    add_body(doc, "PowerShell (tests MohammadNoor barcode 5):")
    add_body(
        doc,
        'Invoke-RestMethod -Method POST -Uri http://127.0.0.1:8080/hikvision/attendance '
        '-ContentType application/json '
        '-Body \'{"subEventType":38,"employeeNoString":"5","currentVerifyMode":"fp",'
        '"status":"success","statusValue":1,"dateTime":"2026-07-07T15:00:00+03:00","serialNo":99999}\''
    )

    add_heading(doc, "Regenerate this PDF", 1)
    add_body(doc, "cd scripts\\hikvision_attendance_service")
    add_body(doc, ".\\.venv\\Scripts\\pip install python-docx docx2pdf")
    add_body(doc, ".\\.venv\\Scripts\\python doc\\generate_documentation.py")
    add_body(doc, f"Output: {DOWNLOADS_PDF}")

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("— End of Document —")
    run.font.color.rgb = COLOR_ACCENT
    run.italic = True
    return doc


def convert_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        return pdf_path.exists()
    except Exception as exc:
        print(f"docx2pdf failed: {exc}")
    try:
        soffice = Path(r"C:\Program Files\LibreOffice\program\soffice.exe")
        if soffice.is_file():
            subprocess.run(
                [str(soffice), "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)],
                check=True,
                timeout=120,
            )
            return pdf_path.exists()
    except Exception as exc:
        print(f"LibreOffice failed: {exc}")
    return False


def main() -> None:
    print("Building simple documentation...")
    doc = build_document()
    doc.save(str(OUTPUT_DOCX))
    print(f"DOCX: {OUTPUT_DOCX}")

    if convert_to_pdf(OUTPUT_DOCX, OUTPUT_PDF):
        shutil.copy2(OUTPUT_PDF, DOWNLOADS_PDF)
        print(f"PDF: {OUTPUT_PDF}")
        print(f"Copied to: {DOWNLOADS_PDF}")
    else:
        print("PDF conversion failed. Open the DOCX in Word and Save As PDF.")


if __name__ == "__main__":
    main()
