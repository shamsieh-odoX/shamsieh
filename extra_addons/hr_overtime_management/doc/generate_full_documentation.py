# -*- coding: utf-8 -*-
"""Generate full HR Overtime Management documentation using company Word template."""

from __future__ import annotations

import shutil
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

TEMPLATE = Path(r"c:\Users\ASUS\Downloads\ملف قالب الشركة (2).docx")
OUTPUT_DOCX = Path(__file__).resolve().parent / "HR_Overtime_Management_Full_Documentation.docx"
OUTPUT_PDF = Path(__file__).resolve().parent / "HR_Overtime_Management_Full_Documentation.pdf"
MODULE_VERSION = "19.0.1.5.3"

COLOR_PRIMARY = RGBColor(0x0E, 0x28, 0x41)
COLOR_ACCENT = RGBColor(0x15, 0x60, 0x82)
COLOR_ACCENT2 = RGBColor(0xE9, 0x71, 0x32)
COLOR_TEXT = RGBColor(0x33, 0x33, 0x33)


def set_paragraph_shading(paragraph, fill_hex: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    p_pr.append(shd)


def add_cover(doc: Document):
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HR Overtime Management System")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = COLOR_PRIMARY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Complete Technical, Configuration & Test Documentation — Odoo 19")
    r.font.size = Pt(16)
    r.font.color.rgb = COLOR_ACCENT

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(
        f"Module: hr_overtime_management (+ optional hr_overtime_payroll)\n"
        f"Version: {MODULE_VERSION}\n"
        f"Document date: {date.today().strftime('%B %d, %Y')}\n"
        f"Database: odoo19 | Config: debian/odoo.conf"
    )
    r.font.size = Pt(11)
    r.font.color.rgb = COLOR_TEXT
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = COLOR_PRIMARY
        set_paragraph_shading(p, "E8E8E8")
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = COLOR_ACCENT
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_ACCENT2
        p.paragraph_format.space_before = Pt(6)
    return p


def add_body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_TEXT
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(f"• {text}")
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_TEXT
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
            set_paragraph_shading(p, "0E2841")
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
            for p in cells[c_idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    doc.add_paragraph()
    return table


def add_toc(doc: Document):
    add_heading(doc, "Table of Contents", 1)
    sections = [
        "1. Executive Summary",
        "2. What Was Customized vs Standard Odoo",
        "3. Business Objectives",
        "4. Delivered Modules & Dependencies",
        "5. System Architecture & Workflow",
        "6. Database Models",
        "7. Cost Calculation",
        "8. Security, Groups & Multi-Company",
        "9. User Interface & Menus",
        "10. Configuration Guide (Step by Step)",
        "11. Server & Database Setup (odoo.conf)",
        "12. End-User Guide",
        "13. Automated Test Cases",
        "14. Issues Fixed During Delivery",
        "15. Technical File Structure",
        "16. Upgrade & Maintenance",
    ]
    for s in sections:
        add_bullet(doc, s)
    doc.add_page_break()


def build_document() -> Document:
    if TEMPLATE.exists():
        shutil.copy2(TEMPLATE, OUTPUT_DOCX)
        doc = Document(str(OUTPUT_DOCX))
        body = doc.element.body
        for child in list(body):
            if child.tag.endswith("p"):
                body.remove(child)
    else:
        doc = Document()

    add_cover(doc)
    add_toc(doc)

    # --- 1 ---
    add_heading(doc, "1. Executive Summary", 1)
    add_body(
        doc,
        "This document describes the complete HR Overtime Management solution built on Odoo 19 "
        "for your organization. Employees submit overtime linked to projects and tasks; requests "
        "flow through department manager → optional upper manager → HR approval; costs are computed "
        "from contract wage and configurable rate multipliers (Regular, Weekend, Day Off); approved "
        "overtime can create timesheet lines and optionally feed payroll."
    )
    add_body(
        doc,
        "The solution follows Odoo best practices: reusable approval engine (similar to Time Off), "
        "per-company configuration, no duplicate payroll tables, scoped sudo() for multi-company reads "
        "(not blanket security bypass), and friendly UI snackbars for access errors."
    )

    # --- 2 ---
    add_heading(doc, "2. What Was Customized vs Standard Odoo", 1)
    add_heading(doc, "2.1 Reused from Standard Odoo", 2)
    add_table(
        doc,
        ["Standard Odoo", "How We Use It"],
        [
            ["hr.employee + parent_id", "Manager / upper-manager approval chain"],
            ["hr.version (contract wage)", "Hourly cost from employee wage"],
            ["project.project / project.task", "Project & task on each request"],
            ["account.analytic.line", "Optional timesheet on HR approval"],
            ["resource.calendar.leaves", "Public holidays for Day Off type"],
            ["mail.thread / mail.activity", "Chatter and approval reminders"],
            ["res.company + company_ids", "Multi-company branches and record rules"],
            ["ir.attachment", "Supporting documents on requests"],
        ],
    )
    add_heading(doc, "2.2 Custom-Built Components", 2)
    add_table(
        doc,
        ["Component", "Purpose"],
        [
            ["hr.overtime.request", "Main overtime record and workflow"],
            ["hr.overtime.type", "Regular / Weekend / Day Off rate categories per branch"],
            ["hr.overtime.approval.line", "Audit trail per approval step"],
            ["hr.approval.chain.service", "Resolves manager → HR chain from employee hierarchy"],
            ["hr.approval.chain.mixin", "Submit, approve, refuse workflow"],
            ["Auto overtime type", "Type picked from date — employee cannot choose manually"],
            ["Per-company type provisioning", "Each branch gets 3 types on company create"],
            ["overtime_error_handler.js", "Snackbar for company/access errors on overtime forms"],
            ["Refuse wizard", "Mandatory reason when refusing"],
            ["QWeb PDF report", "Printable approval document"],
            ["hr_overtime_payroll (optional)", "Glue to payslip inputs"],
        ],
    )
    add_heading(doc, "2.3 Key Customizations Applied", 2)
    for item in [
        "Datetime fields (start_datetime / end_datetime) instead of separate date + float time.",
        "Overtime Type is computed and read-only for employees.",
        "Weekend = Friday & Saturday by default (overtime_weekend_weekdays = 4,5).",
        "Each company branch has independent rate multipliers.",
        "total_cost = hours × hourly_cost × rate_multiplier.",
        "Project domain limited to employee company + shared projects.",
        "Approval lines created via workflow sudo() — employees do not need direct ACL on lines.",
        "Managers see requests via approval_line_ids OR current_approver_id record rules.",
        "company_label Char field avoids direct res.company reads in the UI.",
    ]:
        add_bullet(doc, item)

    # --- 3 ---
    add_heading(doc, "3. Business Objectives", 1)
    for o in [
        "Employees submit overtime with project, task, and attachments.",
        "Multi-level approval mirrors organizational hierarchy + mandatory HR.",
        "Pay rates configurable per branch without code changes.",
        "Approved overtime can post to timesheets automatically.",
        "Cost derived from contract wage, not a parallel engine.",
        "Reusable approval chain for future HR request types.",
        "Multi-company safe: employee company drives record scope.",
    ]:
        add_bullet(doc, o)

    # --- 4 ---
    add_heading(doc, "4. Delivered Modules & Dependencies", 1)
    add_table(
        doc,
        ["Module", "Required", "Purpose", "Depends On"],
        [
            ["hr_overtime_management", "Yes", "Core overtime app", "hr, hr_timesheet, project, mail, resource"],
            ["hr_overtime_payroll", "No", "Payslip overtime input", "hr_overtime_management, hr_payroll"],
        ],
    )

    # --- 5 ---
    add_heading(doc, "5. System Architecture & Workflow", 1)
    add_heading(doc, "5.1 Approval Chain", 2)
    add_bullet(doc, "Step 1: employee.parent_id (or dept manager) → dept_manager")
    add_bullet(doc, "Step 2: parent of dept manager (if different) → upper_manager")
    add_bullet(doc, "Step 3: user in group_overtime_hr_officer → hr (always final)")
    add_body(doc, "Two-stage chain: dept_manager → HR. Three-stage: dept_manager → upper_manager → HR.")

    add_heading(doc, "5.2 Workflow States", 2)
    add_table(
        doc,
        ["State", "Meaning", "Who Can Act"],
        [
            ["draft", "Employee editing", "Employee (owner)"],
            ["submitted", "Waiting first approver", "Dept manager"],
            ["manager_approved", "Dept approved", "Upper manager or HR"],
            ["upper_manager_approved", "Upper approved", "HR officer"],
            ["hr_approved", "Fully approved", "Read-only; timesheet created"],
            ["refused", "Rejected", "Employee may reset to draft"],
            ["cancel", "Cancelled by employee", "Employee"],
        ],
    )

    add_heading(doc, "5.3 Automatic Overtime Type", 2)
    add_table(
        doc,
        ["Category", "When", "Default Multiplier"],
        [
            ["regular", "Sun–Thu (configurable working days)", "1.5×"],
            ["weekend", "Fri & Sat (weekdays 4,5)", "2.0×"],
            ["day_off", "Public holiday on employee calendar", "2.5×"],
        ],
    )
    add_body(doc, "If the period spans multiple categories, the highest multiplier wins.")

    # --- 6 ---
    add_heading(doc, "6. Database Models", 1)
    add_heading(doc, "6.1 hr.overtime.request", 2)
    add_table(
        doc,
        ["Field", "Description"],
        [
            ["name", "Sequence OT/YYYY/MM/00001"],
            ["employee_id", "Employee who worked overtime"],
            ["employee_company_id", "Integer — employee branch (avoids company ACL on compute)"],
            ["company_id", "Stored FK — always employee's company"],
            ["company_label", "Display name via sudo (UI field)"],
            ["start_datetime / end_datetime", "Overtime period"],
            ["overtime_hours", "Computed duration"],
            ["overtime_type_id", "Auto-computed category"],
            ["project_id / task_id", "Timesheet project and task"],
            ["state", "Workflow state"],
            ["current_approver_id", "Active approver user"],
            ["hourly_cost / total_cost", "Computed monetary fields"],
            ["analytic_line_id", "Timesheet line after HR approval"],
            ["attachment_ids", "Supporting documents (company_id=False)"],
        ],
    )
    add_heading(doc, "6.2 hr.overtime.type", 2)
    add_table(
        doc,
        ["Field", "Description"],
        [
            ["category", "regular | weekend | day_off"],
            ["rate_multiplier", "Pay multiplier per branch"],
            ["company_id", "Company branch (one set per company)"],
        ],
    )
    add_heading(doc, "6.3 hr.overtime.approval.line", 2)
    add_table(
        doc,
        ["Field", "Description"],
        [
            ["role", "dept_manager | upper_manager | hr"],
            ["approver_id", "Assigned approver"],
            ["state", "pending | to_approve | approved | refused"],
            ["comment", "Refusal reason"],
        ],
    )
    add_heading(doc, "6.4 res.company Extensions", 2)
    add_table(
        doc,
        ["Setting", "Default", "Purpose"],
        [
            ["overtime_generate_analytic_line", "True", "Create timesheet on HR approval"],
            ["overtime_default_type_id", "Regular", "Working day type"],
            ["overtime_weekend_type_id", "Weekend", "Weekend type"],
            ["overtime_holiday_type_id", "Day Off", "Holiday type"],
            ["overtime_weekend_weekdays", "4,5", "Fri=4, Sat=5"],
            ["overtime_daily_hours_cap", "4.0", "Warning threshold"],
            ["overtime_hours_per_month", "173.33", "Wage ÷ hours fallback"],
        ],
    )

    # --- 7 ---
    add_heading(doc, "7. Cost Calculation", 1)
    add_body(doc, "Formula: total_cost = overtime_hours × hourly_cost × rate_multiplier")
    add_bullet(doc, "Primary hourly: from hr.version wage normalized by work schedule")
    add_bullet(doc, "Fallback: wage ÷ overtime_hours_per_month (default 173.33)")
    add_body(doc, "Example: 2 h Regular at 1.5× with $45.81/h → total ≈ $137.43")

    # --- 8 ---
    add_heading(doc, "8. Security, Groups & Multi-Company", 1)
    add_heading(doc, "8.1 Security Groups", 2)
    add_table(
        doc,
        ["Group", "Users", "Access"],
        [
            ["group_overtime_user", "All internal users", "Create/read own requests; edit draft"],
            ["Managers (record rules)", "Assigned approvers", "Read/write pending requests"],
            ["group_overtime_hr_officer", "HR team", "Read all; final HR approval"],
            ["group_overtime_admin", "Admins", "Full config access"],
        ],
    )
    add_heading(doc, "8.2 Record Rules (hr.overtime.request)", 2)
    add_bullet(doc, "Employee read own: employee_id.user_id = current user")
    add_bullet(doc, "Employee write draft / cancel / reset: state-based domains")
    add_bullet(doc, "Manager: current_approver_id = user OR active approval_line approver")
    add_bullet(doc, "Multi-company: company_id in company_ids")
    add_bullet(doc, "HR officer: read all; write at HR approval stage")

    add_heading(doc, "8.3 Multi-Company Best Practice (Implemented)", 2)
    add_body(
        doc,
        "Odoo's built-in company rule employee restricts res.company reads to the user's "
        "Allowed Companies. The overtime module handles this correctly:"
    )
    for item in [
        "Overtime company_id always equals employee's company — never the session company.",
        "Project dropdown limited to employee company + shared (no company) projects.",
        "Validation raises UserError if project branch not in Allowed Companies.",
        "Scoped sudo() for company settings reads (_sudo_company) and web_read for approvers.",
        "Attachments stored with company_id=False to avoid cross-company errors.",
        "Overtime menus set allowed_company_ids to user.company_ids.",
        "Snackbar (not security bypass) when company AccessError occurs on overtime forms.",
        "Admin fix: add branch under Settings → Users → Allowed Companies if cross-branch needed.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "8.4 Approval Line ACL", 2)
    add_body(
        doc,
        "Employees have read-only ACL on hr.overtime.approval.line. Create/write/unlink "
        "happens inside the workflow engine (hr.approval.chain.mixin) using scoped sudo() "
        "after permission checks (_can_user_approve_line, ownership)."
    )

    # --- 9 ---
    add_heading(doc, "9. User Interface & Menus", 1)
    add_table(
        doc,
        ["Menu", "Who", "Purpose"],
        [
            ["Overtime → My Requests", "Employees", "Own overtime list and form"],
            ["Overtime → To Approve", "Managers / HR", "Pending items assigned to you"],
            ["Overtime → All Requests", "HR officers", "Full list across company"],
            ["Overtime → Configuration → Overtime Types", "Admins", "Rate multipliers per branch"],
            ["Settings → Employees → Overtime", "Admins", "Company overtime settings"],
        ],
    )
    add_body(doc, "Views: List, Form (statusbar, approval tab, chatter), Kanban, Search, Pivot, Graph, PDF report.")

    # --- 10 ---
    add_heading(doc, "10. Configuration Guide (Step by Step)", 1)
    add_heading(doc, "10.1 Initial Installation", 2)
    for i, step in enumerate([
        "Ensure Odoo 19 is running with debian/odoo.conf (see Section 11).",
        "Database name must be odoo19 (not odoo).",
        "Apps → Update Apps List → search HR Overtime Management → Install.",
        "Optional: install hr_overtime_payroll if payroll integration is required.",
        "Upgrade after any code change: -u hr_overtime_management --stop-after-init",
    ], 1):
        add_bullet(doc, f"{i}. {step}")

    add_heading(doc, "10.2 Per Company Branch", 2)
    for i, step in enumerate([
        "Switch to the company branch (top-bar company selector).",
        "Verify 3 overtime types exist (auto-created): Regular, Weekend, Day Off.",
        "Overtime → Configuration → Overtime Types → edit Rate Multiplier per type.",
        "Settings → Employees → Overtime → set weekend weekdays (default 4,5 = Fri/Sat).",
        "Set Daily Hours Warning Cap and Hourly Cost Divisor if needed.",
        "Enable/disable Timesheet Integration (overtime_generate_analytic_line).",
    ], 1):
        add_bullet(doc, f"{i}. {step}")

    add_heading(doc, "10.3 Users & HR Officers", 2)
    for i, step in enumerate([
        "Settings → Users → assign Officer: Overtime HR Approval to HR staff per branch.",
        "Each employee must have an hr.employee linked to their user.",
        "Set Manager (parent_id) on employee for approval chain resolution.",
        "Allowed Companies: add every branch the user must read (e.g. SF + Chicago if cross-branch).",
        "Default Company should match the employee's company.",
    ], 1):
        add_bullet(doc, f"{i}. {step}")

    add_heading(doc, "10.4 Projects & Employees", 2)
    for i, step in enumerate([
        "Projects used for overtime must have Allow Timesheets enabled.",
        "Employee sees only projects from their company or shared projects.",
        "Ensure hr.version has wage set for cost calculation.",
        "Public holidays: Resource → Working Times → Time Off (calendar leaves).",
    ], 1):
        add_bullet(doc, f"{i}. {step}")

    add_heading(doc, "10.5 Test Users (odoo19 database)", 2)
    add_table(
        doc,
        ["User", "Role", "Notes"],
        [
            ["john_doe (id=5)", "Employee", "San Francisco company; submits overtime"],
            ["toto@gmail.com (id=6)", "Dept manager", "Manager of john_doe; approves/refuses"],
            ["admin", "HR / Admin", "HR approval and configuration"],
        ],
    )

    # --- 11 ---
    add_heading(doc, "11. Server & Database Setup (odoo.conf)", 1)
    add_body(doc, "Configuration file: debian/odoo.conf")
    add_table(
        doc,
        ["Setting", "Value", "Notes"],
        [
            ["addons_path", "addons,extra_addons,odoo/addons", "Includes custom modules"],
            ["db_host", "127.0.0.1", "PostgreSQL host"],
            ["db_port", "5432", "PostgreSQL port"],
            ["db_user", "odoo", "Database user"],
            ["http_port", "8069", "Web interface URL"],
            ["Database", "odoo19", "Target database for this project"],
        ],
    )
    add_heading(doc, "11.1 Start & Upgrade Commands", 2)
    add_body(doc, "Start server:")
    add_body(doc, "python odoo-bin -c debian\\odoo.conf -d odoo19")
    add_body(doc, "Upgrade module after code changes:")
    add_body(doc, "python odoo-bin -c debian\\odoo.conf -d odoo19 -u hr_overtime_management --stop-after-init")
    add_body(doc, "Run automated tests:")
    add_body(doc, "python odoo-bin -c debian\\odoo.conf -d odoo19 -u hr_overtime_management --test-tags hr_overtime --stop-after-init")
    add_body(doc, "Regenerate this documentation:")
    add_body(doc, "python extra_addons/hr_overtime_management/doc/generate_full_documentation.py")

    # --- 12 ---
    add_heading(doc, "12. End-User Guide", 1)
    add_heading(doc, "12.1 Employee — Submit Overtime", 2)
    for s in [
        "Overtime → My Requests → New.",
        "Enter start/end datetime, project, task, description.",
        "Overtime Type and cost compute automatically.",
        "Attach files if needed → Save → Submit.",
    ]:
        add_bullet(doc, s)

    add_heading(doc, "12.2 Manager — Approve or Refuse", 2)
    for s in [
        "Overtime → To Approve (only items assigned to you).",
        "Open request → Approve or Refuse.",
        "Refuse requires a reason (wizard).",
        "After refusing, list is empty until a new request is assigned to you.",
    ]:
        add_bullet(doc, s)

    add_heading(doc, "12.3 HR — Final Approval & Config", 2)
    for s in [
        "Users in Officer: Overtime HR Approval perform final approval.",
        "Overtime → All Requests for full visibility.",
        "Configure rates under Overtime → Configuration → Overtime Types.",
    ]:
        add_bullet(doc, s)

    add_heading(doc, "12.4 Company Access Snackbar", 2)
    add_body(
        doc,
        "If you see a warning snackbar about Allowed Companies, either pick a project from "
        "your own branch or ask an administrator to add the missing company under your user profile."
    )

    # --- 13 ---
    add_heading(doc, "13. Automated Test Cases", 1)
    add_body(
        doc,
        "File: tests/test_hr_overtime.py | Tag: hr_overtime | Run with --test-tags hr_overtime"
    )
    add_table(
        doc,
        ["Test Method", "What It Verifies"],
        [
            ["test_overtime_hours_same_day", "Duration 18:00–21:30 = 3.5 hours"],
            ["test_overtime_hours_overnight", "Overnight span = 4.0 hours"],
            ["test_chain_single_manager_two_stages", "Solo employee: dept_manager → HR"],
            ["test_chain_manager_upper_manager_three_stages", "Full 3-step chain"],
            ["test_chain_keeps_hr_when_same_user_as_manager", "HR stage never removed"],
            ["test_refusal_at_stage_one_stops_chain", "Refuse stops workflow; later lines stay pending"],
            ["test_non_approver_cannot_approve", "Intruder gets AccessError"],
            ["test_internal_user_has_overtime_group", "base.group_user implies overtime user"],
            ["test_employee_can_create_draft_without_approve", "Employee creates own draft"],
            ["test_employee_cannot_approve_submitted_request", "Employee cannot self-approve"],
            ["test_cross_company_project_requires_allowed_company", "UserError without branch; OK after Allowed Companies"],
            ["test_employee_sees_overtime_menu", "Overtime menu visible to employees"],
            ["test_auto_overtime_type_weekday_regular", "Sunday → Regular type"],
            ["test_auto_overtime_type_friday_weekend", "Friday → Weekend type"],
            ["test_auto_overtime_type_saturday_weekend", "Saturday → Weekend type"],
            ["test_auto_overtime_type_public_holiday", "Calendar leave → Day Off type"],
            ["test_cost_computation_all_types", "total_cost for Regular, Weekend, Holiday multipliers"],
            ["test_full_approval_flow_three_stages", "Dept → Upper → HR → hr_approved"],
            ["test_full_approval_flow_two_stages", "Solo manager → HR → hr_approved"],
        ],
    )
    add_heading(doc, "13.1 Manual Test Checklist", 2)
    add_table(
        doc,
        ["#", "Scenario", "Expected Result"],
        [
            ["1", "Employee creates draft overtime", "Saved; type auto-set; company = employee branch"],
            ["2", "Employee submits", "State submitted; approval lines created; activity to manager"],
            ["3", "Manager approves", "State advances; next approver gets activity"],
            ["4", "Manager refuses with reason", "State refused; snackbar not raw AccessError"],
            ["5", "HR final approval", "hr_approved; timesheet line if enabled"],
            ["6", "Wrong company project", "UserError or snackbar; not cryptic ACL dialog"],
            ["7", "Multi-company user john_doe", "My Requests shows own records only"],
            ["8", "Manager toto To Approve", "Shows only requests assigned to toto"],
            ["9", "PDF print", "Report generates for approved request"],
            ["10", "Module upgrade", "No missing column errors"],
        ],
    )

    # --- 14 ---
    add_heading(doc, "14. Issues Fixed During Delivery", 1)
    add_table(
        doc,
        ["Issue", "Cause", "Resolution"],
        [
            ["res.company Chicago AccessError", "Session company ≠ employee Allowed Companies", "employee_company_id, allowed_company_ids, scoped sudo, snackbar"],
            ["Cannot create approval.line", "Employees lack create ACL on lines", "Workflow sudo() in approval mixin"],
            ["Manager cannot refuse (read error)", "Record rules + wizard read overtime", "Approver rule via approval_line_ids; wizard sudo"],
            ["My Approvals empty", "Filter used only current_approver_id", "Domain + rule use approval lines; recompute approver"],
            ["Blanket security bypass reverted", "Earlier bypassed res.company rules", "Proper validation + scoped sudo only"],
            ["Submit crash all_user_ids", "Odoo 19 API change", "get_hr_responsible uses all_user_ids"],
            ["Hourly cost $0", "hr.version not hr.contract", "Read wage from version_id"],
            ["Missing DB columns", "Module not upgraded", "SQL migration + -u hr_overtime_management"],
            ["Menu server action TypeError", "Context string not dict", "literal_eval in merge"],
            ["Archived overtime types", "active=False types skipped", "_ensure_overtime_types reactivates"],
            ["web_read recursion", "Missing super() in override", "Fixed then replaced with scoped sudo"],
            ["Attachment company error", "Wrong company_id on files", "company_id=False for overtime attachments"],
        ],
    )

    # --- 15 ---
    add_heading(doc, "15. Technical File Structure", 1)
    add_body(doc, "extra_addons/hr_overtime_management/")
    for f in [
        "models/hr_overtime_request.py — main logic, workflow, multi-company",
        "models/hr_approval_mixin.py — submit/approve/refuse engine",
        "models/hr_approval_chain_service.py — chain resolver",
        "models/hr_overtime_type.py — rate categories",
        "models/hr_overtime_approval_line.py — approval audit",
        "models/res_company.py — per-branch settings",
        "models/ir_attachment.py — attachment company fix",
        "security/security.xml — groups and record rules",
        "security/security_overtime_approver_access.xml — approver rule upgrade hook",
        "static/src/overtime_error_handler.js — snackbar handler",
        "wizard/hr_overtime_refuse_wizard.py — refusal dialog",
        "tests/test_hr_overtime.py — 19 automated tests",
        "hooks.py — post_init: provision overtime types",
        "doc/generate_full_documentation.py — this document generator",
    ]:
        add_bullet(doc, f)

    # --- 16 ---
    add_heading(doc, "16. Upgrade & Maintenance", 1)
    add_body(doc, "Current module version: " + MODULE_VERSION)
    for item in [
        "After any Python/XML change: upgrade hr_overtime_management on odoo19.",
        "Hard-refresh browser (Ctrl+Shift+R) after JS changes.",
        "Use odoo19 database — not odoo.",
        "Do not delete company rule employee — fix Allowed Companies or use scoped sudo.",
        "Regenerate documentation: python doc/generate_full_documentation.py",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of Document —")
    r.font.color.rgb = COLOR_ACCENT
    r.italic = True

    return doc


def convert_to_pdf(docx_path: Path, pdf_path: Path) -> bool:
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        return pdf_path.exists()
    except Exception as e:
        print(f"docx2pdf failed: {e}")
    try:
        import subprocess
        soffice = r"C:\Program Files\LibreOffice\program\soffice.exe"
        subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(docx_path.parent), str(docx_path)],
            check=True,
            timeout=120,
        )
        return pdf_path.exists()
    except Exception as e:
        print(f"LibreOffice failed: {e}")
    return False


def main():
    print("Building DOCX from company template...")
    if not TEMPLATE.exists():
        print(f"WARNING: Template not found at {TEMPLATE}")
    doc = build_document()
    doc.save(str(OUTPUT_DOCX))
    print(f"DOCX saved: {OUTPUT_DOCX}")

    print("Converting to PDF...")
    if convert_to_pdf(OUTPUT_DOCX, OUTPUT_PDF):
        print(f"PDF saved: {OUTPUT_PDF}")
    else:
        print("PDF conversion unavailable. Open DOCX in Word → Save As PDF.")


if __name__ == "__main__":
    main()
